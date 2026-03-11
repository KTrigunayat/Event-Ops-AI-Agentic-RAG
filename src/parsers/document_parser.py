import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from dataclasses import dataclass

@dataclass
class DocumentContent:
    """Represents parsed document content"""
    text: str
    metadata: Dict[str, Any]
    file_type: str
    file_path: str

class DocumentParser:
    """Document parser for PDF and DOCX files"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_document(self, file_path: str) -> DocumentContent:
        """
        Parse document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentContent object with text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                text, metadata = self._parse_pdf(file_path)
            elif file_extension == '.docx':
                text, metadata = self._parse_docx(file_path)
            elif file_extension == '.txt':
                text, metadata = self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_extension}")
            
            # Add common metadata
            metadata.update({
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'file_type': file_extension
            })
            
            self.logger.info(f"Successfully parsed {file_path.name} ({len(text)} characters)")
            
            return DocumentContent(
                text=text,
                metadata=metadata,
                file_type=file_extension,
                file_path=str(file_path)
            )
            
        except Exception as e:
            self.logger.error(f"Error parsing document {file_path}: {e}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse PDF file"""
        text = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', '')),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n\n--- Page {page_num + 1} ---\n\n"
                            text += page_text
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                
        except Exception as e:
            self.logger.error(f"Error reading PDF file {file_path}: {e}")
            raise
        
        return text.strip(), metadata
    
    def _parse_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse DOCX file"""
        text = ""
        metadata = {}
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': str(core_props.revision) if core_props.revision else '',
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                text += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "--- End Table ---\n"
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
        
        return text.strip(), metadata
    
    def _parse_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse TXT file"""
        text = ""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Basic metadata for text files
            metadata = {
                'encoding': 'utf-8',
                'lines': len(text.splitlines()) if text else 0
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                metadata['encoding'] = 'latin-1'
                metadata['lines'] = len(text.splitlines()) if text else 0
            except Exception as e:
                self.logger.error(f"Error reading TXT file {file_path} with alternative encoding: {e}")
                raise
        except Exception as e:
            self.logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
        
        return text.strip(), metadata
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_formats
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return self.supported_formats.copy()
